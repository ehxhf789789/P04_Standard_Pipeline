"""
DOCX Document Parser

Extracts text, tables, images, and metadata from Microsoft Word documents.
Supports ISO/IEC 29500 (OOXML) format.
"""

import time
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional
from datetime import datetime
import uuid
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from .base import (
    DocumentParser,
    DocumentType,
    DocumentMetadata,
    ParseResult,
    Section,
    TableData,
)

logger = logging.getLogger(__name__)


class DOCXParser(DocumentParser):
    """
    Microsoft Word Document Parser

    Features:
    - Text extraction with structure preservation
    - Heading hierarchy detection
    - Table extraction
    - Image extraction
    - Metadata extraction
    - Style analysis
    - Comment extraction
    """

    document_type = DocumentType.DOCX
    supported_extensions = ['.docx', '.doc']

    # Heading style patterns
    HEADING_STYLES = {
        'Heading 1': 1,
        'Heading 2': 2,
        'Heading 3': 3,
        'Heading 4': 4,
        'Heading 5': 5,
        'Heading 6': 6,
        'Title': 0,
    }

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        self.extract_comments = self.config.get('extract_comments', True)
        self.preserve_formatting = self.config.get('preserve_formatting', True)

    async def parse(self, file_path: Path) -> ParseResult:
        """Parse DOCX and extract all content"""
        start_time = time.time()

        result = ParseResult(
            success=False,
            document_type=self.document_type,
            file_path=str(file_path),
            file_hash=self.compute_file_hash(file_path),
        )

        if not HAS_DOCX:
            result.errors.append("python-docx not installed. Run: pip install python-docx")
            return result

        try:
            doc = Document(file_path)

            # Extract metadata
            result.metadata = await self.extract_metadata(file_path)

            # Extract content with structure
            sections, tables, full_text = await self._extract_content(doc)
            result.sections = sections
            result.tables = tables
            result.full_text = full_text

            # Extract comments if enabled
            if self.extract_comments:
                comments = await self._extract_comments(doc)
                result.metadata.custom['comments'] = comments

            # Calculate statistics
            word_count = len(full_text.split())
            result.statistics = {
                'section_count': len(sections),
                'table_count': len(tables),
                'paragraph_count': len(doc.paragraphs),
                'character_count': len(full_text),
                'word_count': word_count,
            }
            result.metadata.word_count = word_count

            result.success = True

        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            result.errors.append(str(e))

        result.processing_time_ms = int((time.time() - start_time) * 1000)
        return result

    async def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract DOCX metadata"""
        metadata = DocumentMetadata()

        if not HAS_DOCX:
            return metadata

        try:
            doc = Document(file_path)
            props = doc.core_properties

            metadata.title = props.title
            metadata.author = props.author
            metadata.creator = props.author
            metadata.subject = props.subject

            # Parse keywords
            if props.keywords:
                metadata.keywords = [k.strip() for k in props.keywords.split(',')]

            # Dates
            if props.created:
                metadata.created = props.created
            if props.modified:
                metadata.modified = props.modified

            # Custom properties
            metadata.custom = {
                'category': props.category,
                'comments': props.comments,
                'content_status': props.content_status,
                'identifier': props.identifier,
                'language': props.language,
                'last_modified_by': props.last_modified_by,
                'revision': props.revision,
                'version': props.version,
            }

            if props.language:
                metadata.language = props.language

        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")

        return metadata

    async def extract_text(self, file_path: Path) -> str:
        """Extract full text from DOCX"""
        if not HAS_DOCX:
            return ""

        try:
            doc = Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)

            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""

    async def _extract_content(
        self, doc: 'Document'
    ) -> tuple[List[Section], List[TableData], str]:
        """Extract structured content from document"""
        sections = []
        tables = []
        text_parts = []

        current_section = Section(
            id=str(uuid.uuid4()),
            heading=None,
            level=0,
        )
        section_stack = [current_section]

        # Track table index for naming
        table_idx = 0

        for element in doc.element.body:
            tag = element.tag.split('}')[-1]  # Remove namespace

            if tag == 'p':  # Paragraph
                para_result = self._process_paragraph(element, doc)

                if para_result['is_heading']:
                    # Save current section if it has content
                    if current_section.paragraphs or current_section.content:
                        sections.append(current_section)

                    # Create new section
                    heading_level = para_result['heading_level']
                    current_section = Section(
                        id=str(uuid.uuid4()),
                        heading=para_result['text'],
                        level=heading_level,
                    )

                    # Update section stack
                    while len(section_stack) > 1 and section_stack[-1].level >= heading_level:
                        section_stack.pop()
                    section_stack.append(current_section)

                else:
                    current_section.paragraphs.append(para_result['text'])
                    current_section.content += para_result['text'] + "\n"

                text_parts.append(para_result['text'])

            elif tag == 'tbl':  # Table
                table_data = self._process_table(element, doc, table_idx)
                tables.append(table_data)
                current_section.tables.append(table_data)
                table_idx += 1

                # Add table text to full text
                for row in table_data.rows:
                    text_parts.append(' | '.join(str(cell) for cell in row))

        # Don't forget the last section
        if current_section.paragraphs or current_section.content:
            sections.append(current_section)

        full_text = "\n".join(text_parts)
        return sections, tables, full_text

    def _process_paragraph(
        self, element: Any, doc: 'Document'
    ) -> Dict[str, Any]:
        """Process a paragraph element"""
        result = {
            'text': '',
            'is_heading': False,
            'heading_level': 0,
            'style': None,
            'alignment': None,
        }

        # Find corresponding paragraph object
        for para in doc.paragraphs:
            if para._element == element:
                result['text'] = para.text
                result['style'] = para.style.name if para.style else None

                # Check if heading
                if para.style and para.style.name in self.HEADING_STYLES:
                    result['is_heading'] = True
                    result['heading_level'] = self.HEADING_STYLES[para.style.name]

                # Check alignment
                if para.alignment:
                    result['alignment'] = str(para.alignment)

                break

        # Fallback: extract text from element
        if not result['text']:
            texts = element.itertext()
            result['text'] = ''.join(texts)

            # Check for outline level (another heading indicator)
            outline_lvl = element.find('.//w:outlineLvl', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if outline_lvl is not None:
                level = int(outline_lvl.get(qn('w:val')))
                result['is_heading'] = True
                result['heading_level'] = level + 1

        return result

    def _process_table(
        self, element: Any, doc: 'Document', index: int
    ) -> TableData:
        """Process a table element"""
        table_data = TableData(
            id=str(uuid.uuid4()),
            name=f"Table_{index + 1}",
        )

        # Find corresponding table object
        for table in doc.tables:
            if table._element == element:
                rows = []

                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())

                    if row_idx == 0:
                        table_data.headers = row_data
                    else:
                        rows.append(row_data)

                table_data.rows = rows
                break

        return table_data

    async def _extract_comments(self, doc: 'Document') -> List[Dict[str, Any]]:
        """Extract comments from document"""
        comments = []

        try:
            # Access comments through XML
            from docx.opc.constants import RELATIONSHIP_TYPE as RT

            comments_part = doc.part.related_parts.get(RT.COMMENTS)
            if comments_part:
                # Parse comments XML
                from lxml import etree
                tree = etree.fromstring(comments_part.blob)

                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

                for comment in tree.findall('.//w:comment', ns):
                    comment_data = {
                        'id': comment.get(qn('w:id')),
                        'author': comment.get(qn('w:author')),
                        'date': comment.get(qn('w:date')),
                        'text': ''.join(comment.itertext()),
                    }
                    comments.append(comment_data)

        except Exception as e:
            logger.warning(f"Could not extract comments: {e}")

        return comments

    async def extract_images(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract images from DOCX"""
        images = []

        if not HAS_DOCX:
            return images

        try:
            doc = Document(file_path)

            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_part = rel.target_part
                    images.append({
                        'rel_id': rel.rId,
                        'target': rel.target_ref,
                        'content_type': image_part.content_type,
                        'size': len(image_part.blob),
                    })

        except Exception as e:
            logger.error(f"Error extracting images: {e}")

        return images

    async def extract_styles(self, file_path: Path) -> Dict[str, Any]:
        """Extract document styles"""
        styles = {}

        if not HAS_DOCX:
            return styles

        try:
            doc = Document(file_path)

            for style in doc.styles:
                if style.type == 1:  # Paragraph style
                    styles[style.name] = {
                        'type': 'paragraph',
                        'base_style': style.base_style.name if style.base_style else None,
                        'font_name': style.font.name if hasattr(style, 'font') else None,
                    }

        except Exception as e:
            logger.error(f"Error extracting styles: {e}")

        return styles
